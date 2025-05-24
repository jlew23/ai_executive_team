"""
Knowledge base implementation using Supabase for vector storage.
This module provides classes for document processing and vector search.
"""

import os
import json
import uuid
import logging
import tempfile
from typing import List, Dict, Any, Optional
import numpy as np
from datetime import datetime
from pathlib import Path

# Import Supabase configuration
from supabase_config import get_supabase_client

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class Document:
    """Document class for knowledge base."""

    def __init__(self, doc_id: str, name: str, content: str, doc_type: str = None, metadata: Dict = None):
        """Initialize a document."""
        self.doc_id = doc_id
        self.name = name
        self.content = content
        self.type = doc_type or self._infer_type(name)
        self.metadata = metadata or {}
        self.embedding = None

    def _infer_type(self, filename: str) -> str:
        """Infer document type from filename."""
        ext = Path(filename).suffix.lower()

        if ext in ['.pdf', '.docx', '.doc']:
            return "Document"
        elif ext in ['.xlsx', '.xls', '.csv']:
            return "Financial Information"
        elif ext in ['.txt', '.md']:
            return "Text"
        elif ext in ['.jpg', '.jpeg', '.png', '.gif']:
            return "Image"
        else:
            return "Other"

    def to_dict(self) -> Dict:
        """Convert document to dictionary."""
        return {
            "doc_id": self.doc_id,
            "name": self.name,
            "type": self.type,
            "content": self.content,
            "uploaded": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "size": f"{len(self.content) / 1024:.1f} KB",
            **self.metadata
        }

    @classmethod
    def from_dict(cls, data: Dict) -> 'Document':
        """Create document from dictionary."""
        metadata = {k: v for k, v in data.items() if k not in ["doc_id", "name", "type", "content"]}
        return cls(
            doc_id=data.get("doc_id", str(uuid.uuid4())),
            name=data.get("name", ""),
            content=data.get("content", ""),
            doc_type=data.get("type"),
            metadata=metadata
        )


class DocumentProcessor:
    """Process documents for the knowledge base."""

    def __init__(self):
        """Initialize document processor."""
        pass

    def process_text(self, text: str, filename: str = "text_input.txt") -> Document:
        """Process text input."""
        doc_id = str(uuid.uuid4())
        return Document(doc_id=doc_id, name=filename, content=text)

    def process_file(self, file_path: str) -> Document:
        """Process file input."""
        try:
            file_path = Path(file_path)
            doc_id = str(uuid.uuid4())

            # Read file content based on extension
            ext = file_path.suffix.lower()

            if ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif ext in ['.pdf']:
                try:
                    # Try to import PyPDF2
                    import PyPDF2

                    # Read PDF content
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        content = ""
                        for page_num in range(len(pdf_reader.pages)):
                            content += pdf_reader.pages[page_num].extract_text() + "\n"

                        if not content.strip():
                            content = f"PDF content extraction failed for {file_path.name}. The PDF may be scanned or have security restrictions."
                except ImportError:
                    # If PyPDF2 is not installed, provide a helpful message
                    content = f"PDF processing requires PyPDF2. Please install it with 'pip install PyPDF2'."
                    logger.error("PyPDF2 not installed. Cannot process PDF files.")
                except Exception as pdf_error:
                    # Handle other PDF processing errors
                    content = f"Error processing PDF {file_path.name}: {str(pdf_error)}"
                    logger.error(f"PDF processing error: {pdf_error}")
            elif ext in ['.docx', '.doc']:
                try:
                    # Try to import docx
                    import docx

                    # Read DOCX content
                    if ext == '.docx':
                        doc = docx.Document(file_path)
                        content = "\n".join([para.text for para in doc.paragraphs])
                    else:  # .doc files not supported by python-docx
                        content = f"DOC format not supported. Please convert to DOCX."
                except ImportError:
                    # If docx is not installed, provide a helpful message
                    content = f"DOCX processing requires python-docx. Please install it with 'pip install python-docx'."
                    logger.error("python-docx not installed. Cannot process DOCX files.")
                except Exception as docx_error:
                    # Handle other DOCX processing errors
                    content = f"Error processing DOCX {file_path.name}: {str(docx_error)}"
                    logger.error(f"DOCX processing error: {docx_error}")
            elif ext in ['.xlsx', '.xls', '.csv']:
                try:
                    # Try to import pandas
                    import pandas as pd

                    # Read spreadsheet content
                    if ext == '.csv':
                        df = pd.read_csv(file_path)
                    else:  # Excel files
                        df = pd.read_excel(file_path)

                    # Convert DataFrame to string representation
                    content = df.to_string(index=False)
                except ImportError:
                    # If pandas is not installed, provide a helpful message
                    content = f"Spreadsheet processing requires pandas. Please install it with 'pip install pandas openpyxl'."
                    logger.error("pandas not installed. Cannot process spreadsheet files.")
                except Exception as excel_error:
                    # Handle other spreadsheet processing errors
                    content = f"Error processing spreadsheet {file_path.name}: {str(excel_error)}"
                    logger.error(f"Spreadsheet processing error: {excel_error}")
            else:
                # For unsupported file types, return an error message as content
                content = f"Unsupported file type: {ext}"
                logger.warning(f"Unsupported file type: {ext}")

            # Create metadata with file information
            metadata = {
                "file_type": ext.lstrip('.').upper(),
                "uploaded_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "size": f"{file_path.stat().st_size / 1024:.1f} KB"
            }

            return Document(doc_id=doc_id, name=file_path.name, content=content, metadata=metadata)

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise


class VectorKnowledgeBase:
    """Vector knowledge base using Supabase."""

    def __init__(self, name="ai_kb", persist_directory=None):
        """Initialize vector knowledge base.

        Args:
            name: Name of the knowledge base
            persist_directory: Directory to persist the knowledge base (not used with Supabase)
        """
        self.supabase = get_supabase_client()
        self.name = name
        self.table_name = "documents"
        self.embedding_dimension = 1536  # Default for OpenAI embeddings
        self.documents = {}  # For compatibility with the old interface

        # Ensure the table exists
        self._ensure_table_exists()

    def _ensure_table_exists(self):
        """Ensure the documents table exists in Supabase."""
        try:
            # Check if table exists by querying it
            self.supabase.table(self.table_name).select("id").limit(1).execute()
            logger.info(f"Table {self.table_name} exists")
        except Exception as e:
            logger.warning(f"Table {self.table_name} does not exist or error: {e}")
            logger.info(f"Creating table {self.table_name}")

            # Create the table with pgvector extension
            # Note: This requires pgvector to be enabled in your Supabase project
            # You may need to enable it in the Supabase dashboard
            try:
                # SQL to create the table with vector support
                sql = f"""
                CREATE EXTENSION IF NOT EXISTS vector;

                CREATE TABLE IF NOT EXISTS {self.table_name} (
                    id UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
                    doc_id TEXT NOT NULL,
                    name TEXT NOT NULL,
                    type TEXT,
                    content TEXT NOT NULL,
                    embedding VECTOR({self.embedding_dimension}),
                    metadata JSONB,
                    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                );

                CREATE INDEX IF NOT EXISTS documents_embedding_idx ON {self.table_name} USING ivfflat (embedding vector_cosine_ops);
                """

                # Execute the SQL
                self.supabase.query(sql).execute()
                logger.info(f"Created table {self.table_name}")
            except Exception as e:
                logger.error(f"Error creating table {self.table_name}: {e}")
                raise

    def add_document(self, document: Document) -> str:
        """Add document to knowledge base."""
        try:
            # Generate embedding for the document
            embedding = self._generate_embedding(document.content)
            document.embedding = embedding

            # Prepare data for insertion
            data = {
                "doc_id": document.doc_id,
                "name": document.name,
                "type": document.type,
                "content": document.content,
                "embedding": embedding.tolist() if embedding is not None else None,
                "metadata": document.metadata
            }

            # Insert into Supabase
            result = self.supabase.table(self.table_name).insert(data).execute()

            logger.info(f"Added document {document.name} to knowledge base")
            return document.doc_id

        except Exception as e:
            logger.error(f"Error adding document to knowledge base: {e}")
            raise

    def search(self, query: str, limit: int = 5, semantic_weight: float = 0.8, keyword_weight: float = 0.2) -> List[Dict]:
        """Search knowledge base for documents matching query."""
        try:
            # Generate embedding for the query
            query_embedding = self._generate_embedding(query)

            if query_embedding is not None:
                # Semantic search using vector similarity
                sql = f"""
                SELECT
                    id,
                    doc_id,
                    name,
                    type,
                    content,
                    metadata,
                    1 - (embedding <=> '{query_embedding.tolist()}') as similarity
                FROM {self.table_name}
                ORDER BY similarity DESC
                LIMIT {limit};
                """

                result = self.supabase.query(sql).execute()

                # Format results
                results = []
                for item in result.data:
                    results.append({
                        "id": item["doc_id"],
                        "text": item["content"],
                        "metadata": {
                            "source_name": item["name"],
                            "type": item["type"]
                        },
                        "score": item["similarity"]
                    })

                return results
            else:
                # Fallback to keyword search
                logger.warning("No embedding generated for query, falling back to keyword search")

                # Simple keyword search
                sql = f"""
                SELECT
                    id,
                    doc_id,
                    name,
                    type,
                    content,
                    metadata
                FROM {self.table_name}
                WHERE content ILIKE '%{query}%'
                LIMIT {limit};
                """

                result = self.supabase.query(sql).execute()

                # Format results
                results = []
                for item in result.data:
                    results.append({
                        "id": item["doc_id"],
                        "text": item["content"],
                        "metadata": {
                            "source_name": item["name"],
                            "type": item["type"]
                        },
                        "score": 0.5  # Default score for keyword search
                    })

                return results

        except Exception as e:
            logger.error(f"Error searching knowledge base: {e}")
            return []

    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get document by ID."""
        try:
            # Try to get document by doc_id first
            result = self.supabase.table(self.table_name).select("*").eq("doc_id", doc_id).execute()

            # If no results, try to get document by id
            if not result.data:
                result = self.supabase.table(self.table_name).select("*").eq("id", doc_id).execute()

            if result.data:
                item = result.data[0]

                # Get doc_id from id if doc_id is not available
                doc_id = item.get("doc_id") or item.get("id")

                # Get name from source_name if name is not available
                name = item.get("name") or item.get("source_name") or "Unknown"

                # Get content from text if content is not available
                content = item.get("content") or item.get("text") or ""

                # Get type from source_type if type is not available
                doc_type = item.get("type") or item.get("source_type") or "Unknown"

                # Get metadata
                metadata = item.get("metadata") or {}

                doc = Document(
                    doc_id=doc_id,
                    name=name,
                    content=content,
                    doc_type=doc_type,
                    metadata=metadata
                )
                return doc

            return None

        except Exception as e:
            logger.error(f"Error getting document {doc_id}: {e}")
            return None

    def delete_document(self, doc_id: str) -> bool:
        """Delete document by ID."""
        try:
            # Try to delete document by doc_id first
            result = self.supabase.table(self.table_name).delete().eq("doc_id", doc_id).execute()

            # If no rows affected, try to delete document by id
            if not result.data or len(result.data) == 0:
                result = self.supabase.table(self.table_name).delete().eq("id", doc_id).execute()

            if result.data and len(result.data) > 0:
                logger.info(f"Deleted document {doc_id}")
                return True
            else:
                logger.warning(f"Document {doc_id} not found for deletion")
                return False

        except Exception as e:
            logger.error(f"Error deleting document {doc_id}: {e}")
            return False

    def list_documents(self) -> List[Dict]:
        """List all documents in knowledge base."""
        try:
            # Select all fields to see what's available
            result = self.supabase.table(self.table_name).select("*").execute()

            documents = []
            # Clear the documents dictionary
            self.documents = {}

            for item in result.data:
                # Log the item for debugging
                logger.info(f"Document from Supabase: {item}")

                # Use id as doc_id if doc_id is not available
                doc_id = item.get("doc_id") or item.get("id")

                # Make sure we have some kind of ID
                if not doc_id:
                    logger.warning(f"Document has no ID: {item}")
                    continue

                # Get name from source_name if name is not available
                name = item.get("name") or item.get("source_name") or "Unknown"

                # Get type from source_type if type is not available
                doc_type = item.get("type") or item.get("source_type") or "Unknown"

                # Get metadata
                metadata = item.get("metadata") or {}

                doc_info = {
                    "doc_id": doc_id,
                    "name": name,
                    "type": doc_type,
                    "source_name": name,
                    "metadata": metadata
                }

                # Add to documents dictionary for compatibility with old interface
                self.documents[doc_id] = doc_info

                # Add to list for return
                documents.append({
                    "doc_id": doc_id,
                    "name": name,
                    "type": doc_type,
                    "uploaded_at": metadata.get("uploaded_at", "Unknown") if metadata else "Unknown",
                    "size": metadata.get("size", "Unknown") if metadata else "Unknown"
                })

            logger.info(f"Retrieved {len(documents)} documents from Supabase")
            for doc in documents:
                logger.info(f"Document: {doc['name']} (ID: {doc['doc_id']})")

            return documents

        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []

    def _generate_embedding(self, text: str) -> Optional[np.ndarray]:
        """Generate embedding for text using OpenAI API or a local model."""
        try:
            # Try to use OpenAI API if available
            try:
                import openai
                import os
                
                # Check if OpenAI API key is available
                api_key = os.getenv("OPENAI_API_KEY")
                if api_key:
                    openai.api_key = api_key
                    
                    # Get embedding from OpenAI
                    response = openai.Embedding.create(
                        input=text,
                        model="text-embedding-ada-002"
                    )
                    
                    # Extract the embedding
                    embedding = np.array(response['data'][0]['embedding'], dtype=np.float32)
                    logger.info(f"Generated embedding using OpenAI API")
                    return embedding
            except (ImportError, Exception) as openai_error:
                logger.warning(f"Could not use OpenAI for embeddings: {openai_error}")
            
            # Try to use a local model if available
            try:
                from sentence_transformers import SentenceTransformer
                
                # Load a sentence transformer model
                model = SentenceTransformer('all-MiniLM-L6-v2')
                
                # Generate embedding
                embedding = model.encode(text, convert_to_numpy=True).astype(np.float32)
                
                # Resize if needed
                if embedding.shape[0] != self.embedding_dimension:
                    logger.warning(f"Resizing embedding from {embedding.shape[0]} to {self.embedding_dimension}")
                    # Pad or truncate to match expected dimension
                    if embedding.shape[0] < self.embedding_dimension:
                        padding = np.zeros(self.embedding_dimension - embedding.shape[0], dtype=np.float32)
                        embedding = np.concatenate([embedding, padding])
                    else:
                        embedding = embedding[:self.embedding_dimension]
                
                logger.info(f"Generated embedding using SentenceTransformer")
                return embedding
            except (ImportError, Exception) as st_error:
                logger.warning(f"Could not use SentenceTransformer for embeddings: {st_error}")
            
            # Fallback to a deterministic but simple embedding method
            # This is better than random but still not great for semantic search
            logger.warning("Falling back to simple hash-based embedding")
            import hashlib
            
            # Generate a deterministic embedding based on text hash
            hash_values = []
            for i in range(self.embedding_dimension):
                # Create different hash for each dimension
                h = hashlib.md5(f"{text}_{i}".encode()).digest()
                # Convert 16 bytes to a float between -1 and 1
                value = int.from_bytes(h[:4], byteorder='little') / (2**32 - 1) * 2 - 1
                hash_values.append(value)
            
            return np.array(hash_values, dtype=np.float32)

        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            # Last resort: random embedding (better than nothing)
            logger.warning("Using random embedding as last resort")
            return np.random.rand(self.embedding_dimension).astype(np.float32)
