"""
Document processing utilities for the knowledge base.
"""

import os
import re
import uuid
import logging
from typing import List, Dict, Any, Optional, Union, BinaryIO, TextIO
from dataclasses import dataclass, field
import requests
from bs4 import BeautifulSoup
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Optional imports for file type support
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """
    A chunk of a document.
    """
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str = ""
    content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class Document:
    """
    A document in the knowledge base.
    """
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    source_type: str = ""  # "file", "text", "url"
    source_name: str = ""
    content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List[DocumentChunk] = field(default_factory=list)
    version: int = 1
    previous_versions: List[Dict[str, Any]] = field(default_factory=list)

class DocumentProcessor:
    """
    Process documents for the knowledge base.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of chunks to split documents into
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    
    def process_file(self, file_path: str) -> Document:
        """
        Process a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Processed document
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.txt':
            return self._process_text_file(file_path)
        elif file_ext == '.pdf' and PDF_SUPPORT:
            return self._process_pdf_file(file_path)
        elif file_ext == '.docx' and DOCX_SUPPORT:
            return self._process_docx_file(file_path)
        elif file_ext == '.csv':
            return self._process_csv_file(file_path)
        elif file_ext in ['.json', '.xml', '.html', '.md']:
            return self._process_text_file(file_path)  # Process as text for now
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def process_text(self, text: str, source_name: str = "direct_input") -> Document:
        """
        Process text.
        
        Args:
            text: Text to process
            source_name: Name of the source
            
        Returns:
            Processed document
        """
        document = Document(
            source_type="text",
            source_name=source_name,
            content=text,
            metadata={
                "source_type": "text",
                "source_name": source_name,
                "char_count": len(text),
                "word_count": len(text.split())
            }
        )
        
        document.chunks = self._split_text(document)
        return document
    
    def process_url(self, url: str) -> Document:
        """
        Process a URL.
        
        Args:
            url: URL to process
            
        Returns:
            Processed document
        """
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "header", "footer", "nav"]):
                script.extract()
            
            # Get text
            text = soup.get_text()
            
            # Clean text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            document = Document(
                source_type="url",
                source_name=url,
                content=text,
                metadata={
                    "source_type": "url",
                    "source_name": url,
                    "title": soup.title.string if soup.title else url,
                    "char_count": len(text),
                    "word_count": len(text.split())
                }
            )
            
            document.chunks = self._split_text(document)
            return document
            
        except Exception as e:
            logger.error(f"Error processing URL {url}: {e}")
            raise
    
    def _process_text_file(self, file_path: str) -> Document:
        """
        Process a text file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Processed document
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        document = Document(
            source_type="file",
            source_name=os.path.basename(file_path),
            content=text,
            metadata={
                "source_type": "file",
                "source_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": "txt",
                "char_count": len(text),
                "word_count": len(text.split())
            }
        )
        
        document.chunks = self._split_text(document)
        return document
    
    def _process_pdf_file(self, file_path: str) -> Document:
        """
        Process a PDF file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Processed document
        """
        if not PDF_SUPPORT:
            raise ImportError("PyPDF2 is required to process PDF files")
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text()
        
        document = Document(
            source_type="file",
            source_name=os.path.basename(file_path),
            content=text,
            metadata={
                "source_type": "file",
                "source_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": "pdf",
                "page_count": len(pdf_reader.pages),
                "char_count": len(text),
                "word_count": len(text.split())
            }
        )
        
        document.chunks = self._split_text(document)
        return document
    
    def _process_docx_file(self, file_path: str) -> Document:
        """
        Process a DOCX file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Processed document
        """
        if not DOCX_SUPPORT:
            raise ImportError("python-docx is required to process DOCX files")
        
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        document = Document(
            source_type="file",
            source_name=os.path.basename(file_path),
            content=text,
            metadata={
                "source_type": "file",
                "source_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "char_count": len(text),
                "word_count": len(text.split())
            }
        )
        
        document.chunks = self._split_text(document)
        return document
    
    def _process_csv_file(self, file_path: str) -> Document:
        """
        Process a CSV file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Processed document
        """
        df = pd.read_csv(file_path)
        text = df.to_string()
        
        document = Document(
            source_type="file",
            source_name=os.path.basename(file_path),
            content=text,
            metadata={
                "source_type": "file",
                "source_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": "csv",
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": list(df.columns),
                "char_count": len(text),
                "word_count": len(text.split())
            }
        )
        
        document.chunks = self._split_text(document)
        return document
    
    def _split_text(self, document: Document) -> List[DocumentChunk]:
        """
        Split text into chunks.
        
        Args:
            document: Document to split
            
        Returns:
            List of document chunks
        """
        text_chunks = self.text_splitter.split_text(document.content)
        
        return [
            DocumentChunk(
                document_id=document.id,
                content=chunk,
                metadata={
                    **document.metadata,
                    "chunk_index": i,
                    "chunk_count": len(text_chunks)
                }
            )
            for i, chunk in enumerate(text_chunks)
        ]
    
    def update_document(self, document_id: str, new_content: str, version_manager: Any) -> Document:
        """
        Update an existing document with new content.
        
        Args:
            document_id: ID of the document to update
            new_content: New content for the document
            version_manager: Version manager to handle versioning
            
        Returns:
            Updated document
        """
        # Get the existing document
        old_document = version_manager.get_document(document_id)
        if not old_document:
            raise ValueError(f"Document with ID {document_id} not found")
        
        # Create a new version of the document
        new_document = Document(
            id=document_id,
            source_type=old_document.source_type,
            source_name=old_document.source_name,
            content=new_content,
            metadata={
                **old_document.metadata,
                "char_count": len(new_content),
                "word_count": len(new_content.split()),
                "updated_at": version_manager.get_current_timestamp()
            },
            version=old_document.version + 1,
            previous_versions=old_document.previous_versions + [{
                "version": old_document.version,
                "content": old_document.content,
                "metadata": old_document.metadata,
                "timestamp": version_manager.get_current_timestamp()
            }]
        )
        
        # Split the new content into chunks
        new_document.chunks = self._split_text(new_document)
        
        return new_document
